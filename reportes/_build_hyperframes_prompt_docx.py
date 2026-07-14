from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reportes/2026-07-10-prompt-analisis-hyperframes-heygen.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(table)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Prompt mejorado: analisis de HeyGen Hyperframes para ensamblado audiovisual")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Courseforge / SofLIA - Engine | Remotion, Cloud Run, Lambda, agentes y automatizacion de video")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_metadata_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    rows = [
        ("Tipo de entrega", "Prompt refinado para solicitar un analisis arquitectonico, operativo y de costos."),
        ("Uso recomendado", "Pegar este prompt en una sesion con acceso al repositorio actual y, si es posible, a documentacion vigente de HeyGen."),
        ("Enfoque", "Diagnostico primero; no implementar hasta tener decision, riesgos, contratos y plan por fases."),
        ("Nota critica", "No asumir capacidades ni precios de Hyperframes: verificar documentacion oficial y pricing actualizado antes de estimar costos."),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), "F2F4F7")
        for cell in table.rows[idx].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            table.cell(idx, 0).paragraphs[0].runs[0].bold = True


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_prompt(doc):
    doc.add_heading("Prompt mejorado", level=1)

    p = doc.add_paragraph()
    p.add_run("Actua como Staff Engineer / Principal Engineer / Software Architect ").bold = True
    p.add_run(
        "con experiencia real en arquitectura de producto, Remotion, video rendering, "
        "automatizacion con agentes, integraciones SaaS de video, costos cloud, seguridad, "
        "observabilidad, bases de datos, pipelines asincronos y mantenibilidad de sistemas production-grade."
    )

    doc.add_heading("Contexto actual", level=2)
    doc.add_paragraph(
        "Estamos trabajando en Courseforge / SofLIA - Engine, una plataforma que genera cursos con IA "
        "y que incluye una fase de produccion visual basada en Remotion, templates, preview, ensamblado "
        "y publicacion. Ya existe un apartado para Remotion y ensamblado, pero actualmente dependemos de "
        "infraestructura cloud para los renders pesados."
    )
    add_bullets(
        doc,
        [
            "Cloud Run puede resolver parte del renderizado, pero puede elevar costos si los videos son largos, si hay concurrencia o si se ejecutan renders con muchos recursos.",
            "Remotion Lambda tiene limites practicos de tiempo, memoria y costo; en nuestro caso, muchas ejecuciones llegan a timeout.",
            "El sistema ya tiene contratos existentes de templates, validacion, versionado, aprobacion, preview/render, estados de produccion, storage y publicacion.",
            "Queremos evaluar si conviene exportar el paso de ensamblado hacia HeyGen Hyperframes para automatizar parte del proceso con ayuda de agentes.",
            "No queremos una respuesta generica sobre HeyGen, Remotion, Lambda o agentes. Queremos un analisis aterrizado al repositorio, a los contratos actuales y a los riesgos reales del producto.",
        ],
    )

    doc.add_heading("Objetivo del analisis", level=2)
    doc.add_paragraph(
        "Analiza, tomando en cuenta todo el repositorio actual, si es conveniente mover o complementar el paso "
        "de ensamblado audiovisual con HeyGen Hyperframes y agentes. La respuesta debe comparar esta opcion "
        "contra mantener Remotion en Cloud Run/Lambda, optimizar el pipeline actual, crear una arquitectura "
        "hibrida o usar Hyperframes solo para ciertos tipos de videos."
    )

    doc.add_heading("Instrucciones obligatorias", level=2)
    add_bullets(
        doc,
        [
            "Antes de recomendar una arquitectura, inspecciona el flujo actual del repositorio: Remotion, templates, preview, render, API, jobs, storage, estados de produccion y publicacion.",
            "No inventes capacidades de HeyGen Hyperframes. Si falta informacion, marca el supuesto y pide verificar documentacion oficial.",
            "Verifica pricing vigente, limites de uso, restricciones de API, modelos de cobro, SLA y terminos relevantes antes de hacer estimaciones de costo.",
            "Diferencia claramente lo que se mantiene en Courseforge, lo que se delegaria a HeyGen y lo que tendria que cambiar en backend, frontend, DB y operaciones.",
            "Prioriza seguridad, trazabilidad, control de versiones, reproducibilidad del resultado, costos reales, experiencia de usuario y deuda tecnica.",
        ],
    )

    doc.add_heading("1. Entendimiento del sistema actual", level=2)
    add_bullets(
        doc,
        [
            "Explica como esta organizado hoy el flujo de produccion visual: Remotion, templates, seleccion de plantilla, preview, ensamblado, render y almacenamiento del resultado.",
            "Identifica los modulos, servicios, tablas, buckets, jobs, actions y endpoints relacionados.",
            "Distingue entre preview interno, render final, templates aprobados, bundles subidos, validacion, versionado y publicacion.",
            "Señala que contratos actuales no se deben romper: tenant/organization_id, estados de produccion, aprobaciones, auditabilidad, almacenamiento y publicacion a SofLIA.",
        ],
    )

    doc.add_heading("2. Diagnostico del problema actual", level=2)
    add_bullets(
        doc,
        [
            "Explica por que Lambda y Cloud Run pueden ser problematicos para renders largos o pesados.",
            "Separa los dominios de fallo: timeout local, timeout Lambda, timeout Cloud Run, preview externo, CodeBuild/build, assets inaccesibles, props invalidas, limites de memoria/CPU y cold starts.",
            "Evalua si el problema principal es costo, tiempo de ejecucion, arquitectura de jobs, tamaño de assets, duracion del video, concurrencia, falta de colas, observabilidad o una combinacion.",
            "Indica que metricas deberiamos medir antes de decidir: duracion promedio de render, duracion P95/P99, costo por minuto, tamaño de assets, tasa de fallo, retries, tiempo de upload/download y costo por curso.",
        ],
    )

    doc.add_heading("3. Evaluacion de HeyGen Hyperframes", level=2)
    add_bullets(
        doc,
        [
            "Describe que capacidades de Hyperframes serian utiles para el caso, separando hechos verificados de supuestos.",
            "Evalua si Hyperframes reemplazaria el ensamblado de Remotion, lo complementaria o solo serviria para ciertos tipos de videos como avatar, talking head, escenas generativas o clips automatizados.",
            "Analiza si los agentes deberian generar instrucciones, preparar assets, seleccionar templates, crear escenas, validar outputs, disparar jobs o solo asistir en la orquestacion.",
            "Determina si Hyperframes puede respetar nuestra necesidad de templates/versionado/aprobacion o si obligaria a aceptar un modelo cerrado con menos control.",
        ],
    )

    doc.add_heading("4. Ventajas esperadas", level=2)
    add_bullets(
        doc,
        [
            "Reduccion potencial de carga en Cloud Run/Lambda si HeyGen asume parte del renderizado pesado.",
            "Automatizacion mas alta con agentes para generar escenas, iterar variantes y preparar videos sin mantener tanto codigo de composicion.",
            "Posible aceleracion de time-to-market si Hyperframes cubre casos visuales que hoy requeririan templates Remotion complejos.",
            "Menor mantenimiento de infraestructura de render si el proveedor abstrae escalado, colas, codecs y procesamiento.",
            "Posible mejora en calidad visual para ciertos formatos si HeyGen ofrece assets, avatares o generacion audiovisual lista para produccion.",
        ],
    )

    doc.add_heading("5. Desventajas y riesgos", level=2)
    add_bullets(
        doc,
        [
            "Dependencia fuerte de proveedor y posible vendor lock-in.",
            "Costos variables que pueden ser dificiles de predecir si el cobro depende de minutos, creditos, concurrencia, resolucion, avatares, voces o generacion por escena.",
            "Menor control sobre el render final comparado con Remotion, especialmente en layouts, timing, versionado exacto y reproducibilidad.",
            "Riesgos de privacidad y compliance al enviar guiones, datos del curso, assets, voces o material de clientes a un proveedor externo.",
            "Integracion compleja si Hyperframes no expone APIs estables para jobs, webhooks, progreso, versionado, plantillas, assets y descarga de resultado.",
            "Dificultad para mantener paridad entre preview web, aprobacion interna, resultado final y publicacion a SofLIA.",
            "Riesgo de que los agentes generen resultados no deterministas o no aprobados si no hay validaciones y gates server-side.",
        ],
    )

    doc.add_heading("6. Alternativas arquitectonicas a comparar", level=2)
    add_numbered(
        doc,
        [
            "Mantener Remotion en Cloud Run y optimizar timeouts, colas, snapshots, assets, caching y observabilidad.",
            "Mantener Remotion como render principal y usar HeyGen solo para generar clips, avatares o segmentos especificos.",
            "Delegar el ensamblado completo a HeyGen Hyperframes, manteniendo Courseforge como orquestador y fuente de verdad.",
            "Arquitectura hibrida: Courseforge decide, versiona y audita; HeyGen produce cuando conviene; Remotion queda como fallback o como motor para plantillas controladas.",
            "Pipeline multi-provider: Remotion local/cloud, HeyGen y otros proveedores compiten segun tipo de video, costo, SLA, calidad y restricciones del tenant.",
        ],
    )
    doc.add_paragraph(
        "Para cada alternativa, compara complejidad, costo, seguridad, mantenibilidad, escalabilidad, experiencia de usuario, impacto en el repo, riesgos, migracion y rollback."
    )

    doc.add_heading("7. Herramientas que se usarian", level=2)
    add_bullets(
        doc,
        [
            "HeyGen Hyperframes, si su API y pricing lo permiten para automatizacion real.",
            "Agentes para planificar escenas, mapear contenido educativo a instrucciones visuales, validar outputs y preparar reintentos.",
            "Backend de orquestacion en Courseforge para crear jobs, firmarlos, guardar estados, recibir webhooks y validar resultados.",
            "Supabase/PostgreSQL para estado, auditoria, organization_id, render jobs, logs, provider, costos estimados/reales, checksums y referencias a assets.",
            "Storage actual para assets de entrada y videos finales, evitando depender solo de URLs temporales de proveedor.",
            "Remotion existente como fallback, motor interno o renderer para plantillas que requieren control exacto.",
            "Sistema de observabilidad: logs estructurados, correlation IDs, metricas por proveedor, errores normalizados y trazabilidad por artifact/lesson/component.",
        ],
    )

    doc.add_heading("8. Herramientas y contratos que se mantienen", level=2)
    add_bullets(
        doc,
        [
            "Autenticacion, roles, permisos y multi-tenancy del sistema actual.",
            "Artifacts, syllabus, instructional plans, materials y material_components como fuente de verdad del contenido.",
            "Validacion, versionado y aprobacion de templates antes de producir videos finales.",
            "Estados de produccion y publicacion a SofLIA.",
            "Storage propio para resultados finales y assets que deban conservarse.",
            "Auditoria de jobs, inputs, outputs, errores, costos y decisiones de agente.",
            "Gates server-side: ningun output generado por agentes o proveedor externo debe aceptarse sin validacion.",
        ],
    )

    doc.add_heading("9. Posible arquitectura recomendada", level=2)
    doc.add_paragraph(
        "Propone una arquitectura hibrida como hipotesis inicial, pero validala contra el repositorio antes de concluir:"
    )
    add_numbered(
        doc,
        [
            "La web conserva UI, seleccion de template, aprobaciones, QA, permisos y visibilidad del progreso.",
            "El backend crea un production_render_job con provider = remotion_cloud, remotion_local, heygen_hyperframes u otro.",
            "Un servicio de orquestacion construye un input snapshot minimo, versionado y auditable.",
            "Si el provider es HeyGen, el backend crea el job externo, sube o firma assets, guarda external_job_id y espera webhooks/progreso.",
            "Los agentes preparan instrucciones de escena, mapean assets y proponen variantes, pero no saltan aprobaciones ni escriben estados finales sin validacion.",
            "Al terminar, el backend descarga o recibe el video final, valida formato/duracion/checksum, lo guarda en storage propio y actualiza estados internos.",
            "Remotion se mantiene como fallback o como renderer principal para templates donde se requiere control exacto.",
        ],
    )

    doc.add_heading("10. Costos", level=2)
    add_bullets(
        doc,
        [
            "Compara costo por curso, costo por minuto final, costo por intento fallido, costo por retry y costo por almacenamiento.",
            "Incluye costos directos de HeyGen, Cloud Run, Lambda, storage, egress, build, colas, observabilidad y soporte operativo.",
            "Incluye costos indirectos: lock-in, debugging, soporte a clientes, fallos de proveedor, cambios de pricing, limites de API y tiempo de desarrollo.",
            "Propone una tabla de decision con escenarios: bajo volumen, volumen medio, alto volumen, videos cortos, videos largos, alto retrabajo, alta personalizacion.",
            "No des cifras definitivas si no verificaste pricing vigente; usa rangos o formulas y explica que datos faltan.",
        ],
    )

    doc.add_heading("11. Seguridad, privacidad y compliance", level=2)
    add_bullets(
        doc,
        [
            "Evalua que datos se enviarian a HeyGen: guiones, imagenes, voces, datos de cliente, metadata de cursos, assets propietarios o contenido sensible.",
            "Propone minimizacion de datos, URLs firmadas temporales, tokens de corta vida, scopes reducidos, webhooks firmados y auditoria de cada llamada externa.",
            "Define validaciones server-side antes de aceptar un video externo: formato, duracion, resolucion, hash, ownership, provider, job_id, estado y relacion con artifact/lesson/component.",
            "Evalua riesgos de prompt injection, instrucciones maliciosas a agentes, outputs no aprobados, filtrado de informacion y abuso de cuotas.",
            "Indica que los agentes no deben tener permisos amplios ni acceso directo a secretos productivos.",
        ],
    )

    doc.add_heading("12. Cambios probables en el repositorio", level=2)
    add_bullets(
        doc,
        [
            "Nuevo dominio o modulo de render providers con interfaz comun: createJob, getStatus, cancelJob, handleWebhook, normalizeError, validateOutput.",
            "Extender production_jobs o crear tabla especifica para render_jobs con provider, external_job_id, input_snapshot, output_asset_id, costs, logs y attempts.",
            "Servicio de costos y diagnostico por provider.",
            "Webhook handler para HeyGen con validacion de firma e idempotencia.",
            "Feature flag por organizacion para habilitar Hyperframes de forma gradual.",
            "UI para seleccionar provider, ver progreso, mostrar errores normalizados y fallback.",
            "Tests de contrato para cada provider, tests de seguridad de webhooks y tests de regresion del flujo Remotion actual.",
        ],
    )

    doc.add_heading("13. Plan por fases", level=2)
    add_numbered(
        doc,
        [
            "Fase 0: medir el pipeline actual de Remotion y documentar timeouts/costos reales.",
            "Fase 1: spike tecnico de HeyGen Hyperframes con un curso de prueba, sin tocar produccion.",
            "Fase 2: crear abstraccion de render providers y normalizacion de errores.",
            "Fase 3: integrar Hyperframes detras de feature flag para un tipo de video acotado.",
            "Fase 4: agregar agentes solo para preparar instrucciones/escenas, no para aprobar ni publicar.",
            "Fase 5: medir calidad, costo, retries, tiempos, soporte y compararlo contra Remotion.",
            "Fase 6: decidir si se amplia, se mantiene como complemento o se descarta.",
        ],
    )

    doc.add_heading("14. Entregables esperados de la respuesta", level=2)
    add_bullets(
        doc,
        [
            "Resumen ejecutivo con recomendacion clara.",
            "Mapa del flujo actual segun el repositorio.",
            "Comparativa de alternativas con ventajas, desventajas y costos.",
            "Arquitectura propuesta con flujo paso a paso.",
            "Riesgos tecnicos, de negocio, seguridad, privacidad y operacion.",
            "Cambios probables en frontend, backend, base de datos, storage, jobs y observabilidad.",
            "Plan incremental de implementacion y validacion.",
            "Criterios de decision: cuando conviene usar HeyGen, cuando conviene Remotion y cuando conviene una arquitectura hibrida.",
        ],
    )

    doc.add_heading("Restricciones finales", level=2)
    add_bullets(
        doc,
        [
            "No propongas una reescritura completa del sistema.",
            "No rompas el flujo actual de Remotion si puede mantenerse como fallback.",
            "No asumas que Hyperframes puede reemplazar todo sin validar API, costos, limites y calidad.",
            "No permitas que agentes aprueben, publiquen o sobrescriban estados criticos sin validacion server-side.",
            "No aceptes outputs externos sin verificarlos, almacenarlos y auditarlos dentro del sistema.",
            "No escondas costos ni lock-in; si faltan datos, declarlos explicitamente.",
        ],
    )


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Prompt arquitectonico - Courseforge / SofLIA - Engine")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main():
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_metadata_table(doc)
    add_prompt(doc)
    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
