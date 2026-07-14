from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reportes" / "2026-07-10-heygen-hyperframes-architecture-analysis.docx"


COLORS = {
    "navy": "1F3A5F",
    "blue": "2563EB",
    "teal": "0F766E",
    "green": "15803D",
    "amber": "B45309",
    "red": "B91C1C",
    "gray": "64748B",
    "light_blue": "EAF2FF",
    "light_teal": "E6FFFB",
    "light_amber": "FFF7ED",
    "light_red": "FEF2F2",
    "light_gray": "F8FAFC",
    "border": "CBD5E1",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = COLORS["border"]) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(COLORS["gray"])


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = COLORS["light_blue"], accent: str = COLORS["blue"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run(body)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], COLORS["navy"])
        set_cell_border(hdr[i], COLORS["navy"])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_border(cells[i])
            if row_idx % 2 == 0:
                set_cell_shading(cells[i], COLORS["light_gray"])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("334155")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Body Text"].font.name = "Aptos"
    styles["Body Text"].font.size = Pt(10)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Aptos Display"
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(10)
    styles["List Number"].font.name = "Aptos"
    styles["List Number"].font.size = Pt(10)


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    add_title(
        doc,
        "Courseforge / SofLIA - Engine",
        "Analisis arquitectonico: Remotion, HeyGen Hyperframes y automatizacion con agentes\n"
        "Fecha: 10 de julio de 2026",
    )

    add_callout(
        doc,
        "Recomendacion ejecutiva",
        "No conviene reemplazar Remotion por HeyGen Hyperframes en este momento. Conviene hacer una arquitectura hibrida y multi-provider: Courseforge mantiene orquestacion, seguridad, versionado, approvals, storage final y publicacion; Remotion sigue como motor principal y fallback; HeyGen Hyperframes se prueba como proveedor complementario para casos acotados donde su modelo HTML/CSS/JS y su pricing por minuto reduzcan tiempo operativo o complejidad visual. Los agentes deben preparar escenas e instrucciones, pero no aprobar, publicar ni mutar estados criticos sin validacion server-side.",
        COLORS["light_teal"],
        COLORS["teal"],
    )

    add_h1(doc, "1. Resumen ejecutivo")
    add_body(
        doc,
        "El sistema actual ya tiene un contrato razonablemente maduro para produccion audiovisual: seleccion de templates, preview interno con Remotion, render final via backend, jobs persistidos, storage propio, estados de material_components y sincronizacion hacia publication_requests para publicar a SofLIA. Ese contrato no deberia romperse para integrar HeyGen.",
    )
    add_body(
        doc,
        "HeyGen Hyperframes, segun documentacion oficial verificada el 10 de julio de 2026, permite renderizar composiciones HTML/CSS/JS a video por API, con variables, polling, webhooks y billing por minuto de salida. Eso lo vuelve interesante para composiciones generativas, escenas explicativas, talking-head/avatar + motion graphics, y automatizacion por agentes. Sin embargo, no hay evidencia de que ejecute directamente templates actuales de Remotion; por tanto, usarlo como reemplazo total implicaria reimplementar o adaptar templates y cambiar el modelo de preview.",
    )
    add_body(
        doc,
        "La decision de menor riesgo es construir una abstraccion de render providers y habilitar HeyGen detras de feature flags por organizacion y tipo de video. El MVP debe probar un solo tipo de componente, guardar trazabilidad completa en production_jobs, descargar el resultado a storage propio, validar servidor-side y requerir QA humano antes de publicacion.",
    )

    add_h1(doc, "2. Base de analisis")
    add_h2(doc, "Repositorio inspeccionado")
    add_bullets(
        doc,
        [
            "Backend Remotion: apps/api/src/features/production/remotion-worker.service.ts, remotion-render-orchestrator.service.ts, remotion-render.config.ts, remotion-queue.service.ts, remotion-lambda.provider.ts, remotion-lambda-progress.service.ts.",
            "Contratos de provider: apps/api/src/features/production/render-provider.types.ts, actualmente con provider local y lambda.",
            "Controlador API: apps/api/src/features/production/production.controller.ts.",
            "Acciones web: apps/web/src/domains/materials/actions/production.actions.ts.",
            "UI de postproduccion: apps/web/src/domains/materials/components/PostproductionAssemblyContainer.tsx y players de preview Remotion.",
            "Templates: apps/web/src/domains/production/actions/templates.actions.ts y validation/bundle-validator.ts.",
            "Base de datos: migraciones de production_jobs, production_assets, remotion_template_versions, remotion_template_builds, bucket production-videos y publication_requests.",
        ],
    )
    add_h2(doc, "Fuentes externas verificadas")
    add_body(
        doc,
        "Se consulto documentacion oficial de HeyGen para evitar inventar capacidades: Developer Overview, Quick Start, Hyperframes Overview, Hyperframes API, pricing, usage limits y webhooks. Los precios y limites externos deben volver a verificarse antes de una decision contractual, porque pueden cambiar sin que el repositorio lo refleje.",
    )

    add_h1(doc, "3. Entendimiento del sistema actual")
    add_h2(doc, "Flujo de produccion visual observado")
    add_numbered(
        doc,
        [
            "El admin trabaja desde la UI de postproduccion, donde selecciona template, revisa preview y dispara ensamblado.",
            "El frontend construye o normaliza assets del componente, valida que haya insumos renderizables y llama al backend Express mediante /api/v1/production/remotion/render.",
            "El backend crea o usa un production_job con organization_id, artifact_id, component_id, job_type, provider, status, input_snapshot y datos de idempotencia.",
            "El orquestador decide provider actual entre local y lambda usando configuracion de entorno. Remotion local usa Chromium/FFmpeg del entorno, Lambda usa el provider remoto.",
            "El worker carga material_components, remotion_templates y contexto relacionado; genera props, hashes y diagnostics; renderiza la composicion Remotion interna; sube el MP4 a Supabase Storage en production-videos.",
            "Al completar, actualiza production_jobs, material_components.assets.final_video_url, production_status y sincroniza publication_requests.lesson_videos para que la publicacion a SofLIA consuma el contrato interno.",
        ],
    )
    add_h2(doc, "Diferencias relevantes")
    add_table(
        doc,
        ["Area", "Comportamiento actual", "Implicacion para HeyGen"],
        [
            ["Preview interno", "Remotion Player usa props normalizadas y composiciones internas.", "No debe prometer fidelidad si HeyGen genera un resultado distinto; se requeriria preview especifico o thumbnail/proxy externo."],
            ["Render final", "Backend ejecuta Remotion y controla upload final.", "HeyGen debe ser un provider, no la fuente final de verdad; el resultado debe volver a storage propio."],
            ["Templates aprobados", "Existen estados, versionado, builds y aprobacion manual.", "Solo se deben enviar a HeyGen templates/estilos aprobados o transformaciones versionadas."],
            ["Bundles subidos", "Validacion ZIP estatica y builds controlados.", "Hyperframes usa ZIP HTML/CSS/JS; requiere contrato propio o adaptador, no ejecucion libre de bundles de usuario."],
            ["Produccion", "Estados en production_jobs y material_components.", "Debe extenderse provider/status/error/cost, manteniendo compatibilidad."],
            ["Publicacion", "SofLIA lee lesson_videos y storage interno.", "No debe depender de URLs efimeras del proveedor."],
        ],
        [1.4, 2.3, 2.7],
    )
    add_h2(doc, "Contratos que no deben romperse")
    add_bullets(
        doc,
        [
            "Multi-tenancy: organization_id debe viajar en job, assets, templates, permisos y storage.",
            "Roles y permisos: admins/builders/architects no deben saltarse approval ni RLS por usar proveedor externo.",
            "Estados de produccion: PENDING, IN_PROGRESS, COMPLETED/FAILED y jobs deben seguir auditables.",
            "Aprobaciones manuales: templates, outputs y publicacion no deben quedar en manos de agentes ni webhooks externos.",
            "Storage propio: production-videos sigue siendo la fuente estable para publicacion y auditoria.",
            "Validacion server-side: props, assets, checksums, duracion, formato, resolucion y ownership se validan antes de aceptar el resultado.",
            "Publicacion a SofLIA: publication_requests.lesson_videos conserva su contrato interno.",
        ],
    )

    add_h1(doc, "4. Diagnostico del problema actual")
    add_body(
        doc,
        "El problema no es solamente Cloud Run o Lambda. Es una combinacion de duracion de videos, tamano de assets, estrategia de jobs, limites de infraestructura, observabilidad incompleta para costos por intento y dependencia de Chromium/FFmpeg. Subir timeouts puede aliviar sintomas, pero no reduce el costo por fallo ni mejora la trazabilidad del proveedor.",
    )
    add_table(
        doc,
        ["Dominio de fallo", "Riesgo", "Lectura arquitectonica"],
        [
            ["Timeout local", "Render pesado supera LOCAL_RENDER_TIMEOUT o recursos de la maquina.", "Necesita metricas de duracion, FPS efectivo, uso de CPU/RAM y tamano de assets."],
            ["Timeout Lambda", "Lambda tiene limite practico cercano a 15 minutos y costo por memoria/tiempo.", "Adecuado para clips cortos paralelizables, riesgoso para ensamblados largos."],
            ["Timeout Cloud Run", "Puede admitir mas tiempo, pero cobra CPU/memoria y escala con concurrencia.", "Resuelve limite temporal, no necesariamente costo o debugging."],
            ["Preview externo", "CORS/CORP, URLs firmadas, media inaccesible.", "Debe medirse separado del render final."],
            ["Build templates", "Bundles externos requieren validacion, build y aprobacion.", "No deben mezclarse fallos de template con fallos de render."],
            ["Assets inaccesibles", "URLs expiradas, storage privado, egress lento.", "Afecta Remotion y HeyGen; se requiere asset manifest con TTL y checksum."],
            ["Props invalidas", "Inputs no compatibles con template/composicion.", "Contrato Zod/JSON versionado antes de crear job externo."],
            ["Cold starts", "Chromium/FFmpeg y bundle tardan en iniciar.", "Importa para UX y costo en baja frecuencia."],
            ["Retries", "Reintentos replican costo y pueden duplicar outputs.", "Idempotencia y attempts por provider son obligatorios."],
        ],
        [1.5, 2.1, 2.8],
    )
    add_h2(doc, "Metricas que deben medirse antes de decidir")
    add_bullets(
        doc,
        [
            "Duracion de render promedio, P95 y P99 por tipo de video y template.",
            "Costo por minuto renderizado, costo por video final y costo por curso.",
            "Tasa de fallo por provider, tipo de error y etapa: preview, build, render, upload, sync de publicacion.",
            "Numero de retries y costo de intentos fallidos.",
            "Tamano promedio y P95 de assets; tiempo de download/upload y egress.",
            "Tiempo visible para usuario desde click hasta preview/final.",
            "CPU/RAM usados por Remotion y duracion final del video.",
            "Costo por proveedor: Remotion local, Lambda, Cloud Run, Hyperframes, avatares/agentes.",
        ],
    )

    add_h1(doc, "5. Evaluacion de HeyGen Hyperframes")
    add_h2(doc, "Hechos verificados")
    add_bullets(
        doc,
        [
            "Hyperframes renderiza una composicion HTML/CSS/JS autocontenida a video, usando un navegador headless y exportando MP4/MOV/WebM/GIF/secuencia PNG segun documentacion oficial.",
            "La API permite crear renders con ZIP por URL, asset_id o base64, definir fps, calidad, formato, resolucion, aspect ratio, variables, callback_url y callback_id.",
            "El estado del render se consulta por render_id con estados como queued, rendering, completed y failed; tambien existen webhooks para success/fail.",
            "El pricing oficial consultado indica billing por minuto de output para Hyperframes: 1080p/30 a USD 0.10/min, 1080p/60 a USD 0.20/min, 4K/30 a USD 0.15/min y 4K/60 a USD 0.30/min.",
            "HeyGen tambien ofrece APIs de Video Agent, avatares, lipsync y TTS con otros precios por segundo; no son equivalentes a Hyperframes y deben costearse por separado.",
            "Los limites oficiales incluyen concurrencia/rate limits, tamano de recursos, requisitos de URLs publicas y maximos de duracion/escenas para ciertos endpoints.",
        ],
    )
    add_h2(doc, "Supuestos y puntos a validar")
    add_bullets(
        doc,
        [
            "No se debe asumir que Hyperframes ejecuta templates Remotion existentes. Lo mas probable es que requiera un adaptador o templates Hyperframes separados basados en HTML/CSS/JS.",
            "No se debe asumir que el preview de Remotion sera fiel al output de HeyGen. Se requiere preview especifico, render draft o comparacion visual.",
            "SLA, retencion de datos, terminos de uso, propiedad de outputs, uso de contenido para entrenamiento y DPA/compliance deben validarse comercialmente.",
            "La calidad, duracion maxima efectiva, latencia real, limites de concurrencia por plan y comportamiento ante fallos deben probarse con un spike.",
            "La integracion con agentes debe limitarse a preparacion de escenas/instrucciones y validacion asistida; no debe aprobar ni publicar.",
        ],
    )
    add_h2(doc, "Encaje con Courseforge")
    add_body(
        doc,
        "Hyperframes encaja mejor como provider complementario para escenas visuales parametrizables, motion graphics y videos donde el output HTML/CSS/JS pueda generarse de forma controlada. Es menos adecuado como reemplazo total inmediato porque Courseforge ya invirtio en contratos Remotion, templates aprobados, preview interno y sincronizacion con production_jobs/material_components/publication_requests.",
    )

    add_h1(doc, "6. Ventajas y desventajas")
    add_table(
        doc,
        ["Ventajas esperadas", "Desventajas y riesgos"],
        [
            ["Reduce carga de Cloud Run/Lambda para ciertos videos.", "Vendor lock-in y dependencia de disponibilidad, pricing y limites externos."],
            ["Billing por minuto de salida puede ser predecible para Hyperframes.", "Costos variables en Video Agent, avatares, retries, egress y soporte."],
            ["Menos mantenimiento de Chromium/FFmpeg en infraestructura propia para esos casos.", "Menor control sobre layout, timing, reproducibilidad y debugging."],
            ["Los agentes pueden generar instrucciones/variables para escenas.", "Riesgo de prompt injection, instrucciones no aprobadas y sobreconsumo si se dan permisos amplios."],
            ["Puede facilitar variantes visuales y talking-head/avatar si aplica.", "Privacidad: se enviarian guiones, assets y posiblemente datos de cliente al proveedor."],
            ["Webhooks/polling permiten integracion asincronica.", "Preview fiel y QA pueden volverse mas complejos que con Remotion local."],
        ],
        [3.0, 3.0],
    )

    add_h1(doc, "7. Alternativas arquitectonicas")
    add_table(
        doc,
        ["Alternativa", "Evaluacion", "Cuándo conviene", "Riesgo principal"],
        [
            ["1. Remotion Cloud Run optimizado", "Complejidad media, bajo cambio de repo, mantiene control y reproducibilidad. Costo puede crecer con videos largos/concurrencia.", "Cuando la prioridad es continuidad y control sobre templates existentes.", "Seguir pagando por renders lentos/fallidos si no se mejora observabilidad y colas."],
            ["2. Remotion principal + HeyGen clips", "Complejidad media-alta, buen rollback, conserva Remotion como ensamblador y usa HeyGen para clips/avatar/escenas.", "Cuando ciertos clips se benefician de proveedor generativo sin mover todo el pipeline.", "Integracion de assets y consistencia visual entre proveedores."],
            ["3. Delegar todo a Hyperframes", "Complejidad alta y cambio fuerte de templates/preview. Potencial menor operacion propia, pero alto lock-in.", "Solo si spike demuestra calidad, costos, SLA y compatibilidad superior en la mayoria de casos.", "Reescritura encubierta y perdida de control del pipeline actual."],
            ["4. Hibrida con fallback", "Complejidad alta pero incremental. Courseforge orquesta, decide provider y valida outputs. Mejor balance riesgo/control.", "Opcion recomendada para MVP y escalamiento gradual.", "Necesita contrato de providers, metricas y operacion disciplinada."],
            ["5. Multi-provider", "Mas escalable y flexible; permite costo/calidad/SLA por tipo de video. Mayor complejidad operativa.", "Cuando haya volumen suficiente y varios casos de uso diferenciados.", "Deuda tecnica si se implementa sin contrato comun y normalizacion de errores."],
            ["6. Agentes solo para preparacion", "Bajo riesgo si no mutan estados criticos. Mejora productividad sin ceder control.", "Siempre como regla de seguridad inicial.", "Falsa confianza en outputs de agentes si no hay validaciones."],
        ],
        [1.25, 2.2, 1.75, 1.8],
    )

    add_h1(doc, "8. Herramientas y responsabilidades")
    add_table(
        doc,
        ["Herramienta", "Responsabilidad recomendada", "No debe hacer"],
        [
            ["Courseforge backend", "Orquestar jobs, preparar input_snapshot, elegir provider, validar outputs, actualizar estados, auditar.", "Delegar ownership final de estado o publicacion a un proveedor externo."],
            ["Supabase/PostgreSQL", "Persistir production_jobs, attempts, logs, checksums, costs, provider, organization_id y auditoria.", "Guardar secretos de proveedor sin cifrado o sin scopes."],
            ["Supabase Storage production-videos", "Guardar resultado final aceptado y servir contrato interno.", "Depender de video_url efimera de HeyGen para publicar."],
            ["Remotion", "Provider principal/fallback para templates actuales, preview fiel y renders controlados.", "Cargar bundles no aprobados o saltar validacion."],
            ["HeyGen Hyperframes", "Provider opcional para composiciones HTML/CSS/JS aprobadas y parametrizadas.", "Reemplazar approvals, storage, validacion o tenant isolation."],
            ["HeyGen Avatar/Video Agent", "Opcional para clips concretos de avatar/talking head, costeado por separado.", "Tomar decisiones de contenido, aprobar QA o publicar."],
            ["Agentes", "Preparar escenas, variables, instrucciones, QA asistido y resumen de logs.", "Tener secretos productivos, aprobar estados criticos o ejecutar uploads sin backend."],
            ["Webhooks", "Actualizar progreso externo con firma e idempotencia.", "Aceptar cambios sin validacion de firma, tenant y job ownership."],
            ["Feature flags", "Habilitar provider por organizacion/tipo de video.", "Activar rollout global sin medicion."],
        ],
        [1.55, 2.5, 2.4],
    )

    add_h1(doc, "9. Arquitectura propuesta")
    add_callout(
        doc,
        "Arquitectura recomendada",
        "Hibrida y multi-provider. Courseforge conserva orquestacion, estados, storage final, permisos, versionado y publicacion. Remotion permanece como provider principal y fallback. HeyGen Hyperframes se incorpora como provider experimental por feature flag, inicialmente para un solo tipo de video o clip.",
        COLORS["light_blue"],
        COLORS["blue"],
    )
    add_h2(doc, "Flujo propuesto")
    add_numbered(
        doc,
        [
            "Admin selecciona o confirma template/estilo de produccion en Courseforge.",
            "Backend crea production_render_job con organization_id, artifact_id, component_id, template_version_id, idempotency_key e input_snapshot minimo.",
            "Policy de provider decide entre remotion-cloud, remotion-local, hyperframes o fallback usando feature flags, tipo de video, duracion estimada, costos y SLA.",
            "Si el provider es HeyGen, backend prepara ZIP/asset aprobado, variables, URLs firmadas temporales, callback_url y metadata no sensible.",
            "Agentes pueden ayudar a construir escenas o variables, pero el backend valida schema, aprobacion del template y limites antes de enviar.",
            "Backend envia el job a HeyGen y guarda external_job_id/render_id, request hash, provider, estimated_cost y attempt.",
            "Progreso llega por webhook firmado o polling; cada actualizacion se valida por firma, idempotencia, tenant y estado permitido.",
            "Al completar, backend descarga el video, calcula checksum, valida formato, duracion, resolucion, ownership y relacion con artifact/lesson/component.",
            "Resultado aceptado se sube a production-videos; se actualiza material_components.assets.final_video_url, production_status y production_jobs.",
            "Admin revisa y aprueba. Publication_requests.lesson_videos sigue usando el contrato interno y SofLIA no depende del proveedor externo.",
        ],
    )
    add_h2(doc, "Diagrama textual")
    add_code_block(
        doc,
        """Admin UI
  -> production.actions.ts
  -> Express /api/v1/production/render
  -> production_jobs(input_snapshot, tenant, template_version, idempotency)
  -> ProviderPolicy(feature flags, duration, template, org)
      -> Remotion Provider(local/lambda/cloud) -> production-videos
      -> HeyGen Hyperframes Provider -> HeyGen API -> webhook/polling -> download -> validate -> production-videos
  -> material_components.assets.final_video_url
  -> publication_requests.lesson_videos
  -> QA humano
  -> Publicacion a SofLIA""",
    )

    add_h1(doc, "10. Costos")
    add_body(
        doc,
        "No se deben fijar presupuestos definitivos sin confirmar pricing, SLA, terminos, limites de concurrencia, retencion de datos y moneda/facturacion vigentes. Las cifras de HeyGen indicadas son referencias oficiales consultadas el 10 de julio de 2026 y deben tratarse como inputs verificables, no como contrato.",
    )
    add_code_block(
        doc,
        """costo_hyperframes = minutos_finales * tarifa_hyperframes(resolucion, fps)
costo_avatar = segundos_avatar * tarifa_avatar_modelo
costo_video_agent = segundos_output * tarifa_video_agent
costo_remotion_cloud = tiempo_cpu_memoria * tarifa_cloud + storage + egress
costo_total_job = costo_provider + costo_retries + storage + egress + agentes + QA + operacion
costo_por_curso = sum(costo_total_job por video) + fallos + publicacion + soporte""",
    )
    add_table(
        doc,
        ["Escenario", "Lectura de costo", "Decision sugerida"],
        [
            ["Bajo volumen", "Cloud Run optimizado puede ser suficiente; Hyperframes sirve para comparar calidad/tiempo.", "No migrar. Hacer spike controlado."],
            ["Volumen medio", "Retries y tiempos largos empiezan a pesar; billing por minuto de Hyperframes puede ser atractivo para escenas compatibles.", "Arquitectura hibrida con metricas por provider."],
            ["Alto volumen", "Lock-in y concurrencia importan tanto como precio unitario.", "Multi-provider con contratos, alertas, budgets y negociacion enterprise."],
            ["Videos cortos", "Lambda/Remotion puede seguir siendo competitivo.", "Usar Remotion si el preview y output son confiables."],
            ["Videos largos", "Timeouts y retries vuelven caro el render propio.", "Evaluar Hyperframes o Cloud Run largo, pero medir egress y QA."],
            ["Alta personalizacion", "Agentes y generacion externa pueden aumentar variabilidad y costo.", "Limitar a templates aprobados y variables validadas."],
        ],
        [1.4, 2.55, 2.05],
    )

    add_h1(doc, "11. Seguridad, privacidad y compliance")
    add_h2(doc, "Riesgos al enviar datos a HeyGen")
    add_bullets(
        doc,
        [
            "Guiones, imagenes, audio, voces, material educativo, metadata de cursos y assets propietarios podrian salir del perimetro de Courseforge.",
            "URLs firmadas mal configuradas pueden exponer assets entre tenants o por mas tiempo del necesario.",
            "Webhooks no verificados pueden manipular estados de production_jobs.",
            "Outputs externos pueden tener formato, duracion, ownership o contenido distinto al esperado.",
            "Prompts e instrucciones de agentes pueden filtrar informacion o generar contenido no aprobado.",
        ],
    )
    add_h2(doc, "Controles recomendados")
    add_bullets(
        doc,
        [
            "Minimizacion de datos: enviar solo assets y variables necesarias, sin secretos ni datos de cliente innecesarios.",
            "URLs firmadas temporales con TTL corto, scope reducido y checksums esperados.",
            "API keys en backend, nunca en desktop/browser ni en prompts de agentes.",
            "Webhooks con firma, rotacion de secret, idempotency_key, replay protection y verificacion de ownership.",
            "Templates y ZIPs aprobados solamente; validacion de rutas, tamano, extension, dependencias y variables.",
            "Validacion server-side del output antes de aceptar: MIME, codec, resolucion, duracion, checksum, relacion con job y tenant.",
            "Logs sin secretos, con correlation_id, provider, external_job_id, attempt y error normalizado.",
            "Feature flags por organizacion y budget limits por proveedor.",
            "QA humano antes de publicacion a SofLIA.",
            "Agentes con permisos minimos: pueden proponer, no aprobar; pueden preparar instrucciones, no mutar estados criticos.",
        ],
    )

    add_h1(doc, "12. Cambios probables en el repositorio")
    add_table(
        doc,
        ["Cambio", "Ubicacion probable", "Objetivo"],
        [
            ["Abstraccion de providers", "apps/api/src/features/production/render-provider.types.ts y orchestrator", "Agregar provider hyperframes/heygen sin romper local/lambda."],
            ["Provider HeyGen", "apps/api/src/features/production/providers/heygen-hyperframes.provider.ts", "Crear render, polling/webhook, normalizar errores y descargar output."],
            ["Provider policy", "apps/api/src/features/production/provider-policy.service.ts", "Elegir provider por org, feature flag, tipo de video, duracion y costo."],
            ["Extender production_jobs/attempts", "supabase/migrations", "Guardar provider, external_job_id, attempt, costs, logs, input/output checksums y fallback."],
            ["Webhook handler", "apps/api/src/features/production/production.controller.ts", "Recibir success/fail, validar firma e idempotencia."],
            ["Feature flags por org", "DB/config", "Rollout controlado y rollback rapido."],
            ["UI de progreso/provider", "PostproductionAssemblyContainer.tsx", "Mostrar provider usado, progreso, fallos y fallback."],
            ["Metricas y costos", "jobs/logging/observability", "Comparar tiempo, costo, retry y calidad por provider."],
            ["Tests de contrato", "apps/api tests", "Asegurar que Remotion fallback y HeyGen provider respetan estados."],
            ["Docs de rollback", "docs/ o reportes/", "Operacion segura ante fallos del proveedor externo."],
        ],
        [1.6, 2.2, 2.2],
    )

    add_h1(doc, "13. Plan de implementacion por fases")
    add_table(
        doc,
        ["Fase", "Objetivo", "Criterio de salida"],
        [
            ["0. Diagnostico", "Medir renders actuales, timeouts, costos, retries, tamano de assets y tipos de videos.", "Dashboard o reporte con P50/P95/P99, costo estimado y top fallos."],
            ["1. Spike tecnico", "Probar Hyperframes con un caso pequeno y datos no sensibles.", "API validada, output descargado, costo/latencia/calidad medidos, limites entendidos."],
            ["2. Providers", "Crear contrato interno de providers manteniendo Remotion actual.", "Remotion funciona igual; tests de contrato pasan; espacio listo para HeyGen."],
            ["3. Integracion controlada", "Agregar HeyGen detras de feature flag para un tipo de video.", "Jobs, logs, costos, webhooks/polling y validacion server-side funcionando."],
            ["4. Agentes", "Usar agentes solo para preparar escenas/variables y QA asistido.", "No tienen permisos de aprobacion/publicacion; outputs pasan validadores."],
            ["5. Comparacion real", "Comparar costo, tiempo, calidad, retries y soporte con Remotion.", "Decision informada: ampliar, mantener complemento o descartar."],
            ["6. Escalamiento", "Habilitar por organizacion, fallback y alertas.", "Runbook, budgets, metricas, rollback y soporte documentados."],
        ],
        [1.3, 3.0, 2.1],
    )

    add_h1(doc, "14. Recomendacion final")
    add_callout(
        doc,
        "Decision accionable",
        "Usar HeyGen Hyperframes si, pero como complemento experimental y provider opcional, no como reemplazo total de Remotion. El MVP mas seguro es un provider Hyperframes detras de feature flag, activado para un tipo de video no critico, con Remotion como fallback, storage final en Courseforge, validacion server-side y QA humano.",
        COLORS["light_amber"],
        COLORS["amber"],
    )
    add_bullets(
        doc,
        [
            "Lo que reduce mas riesgo: arquitectura hibrida con Remotion fallback y contratos internos intactos.",
            "Lo que reduce mas costo a corto plazo: diagnostico de renders, normalizacion de retries y uso selectivo de Hyperframes donde el costo por minuto sea menor que render propio fallido/lento.",
            "Lo que mantiene mas control: Courseforge como source of truth para jobs, templates aprobados, storage, auditoria y publicacion.",
            "Lo que escala mejor: multi-provider con metricas, budgets, feature flags, normalizacion de errores y fallback.",
            "Lo que no deberiamos hacer: migrar todos los templates a Hyperframes sin spike, entregar secretos a agentes, depender de URLs externas para publicar, aceptar outputs sin validacion o permitir que un webhook externo cambie estados criticos sin firma e idempotencia.",
        ],
    )

    add_h1(doc, "15. Fuentes oficiales consultadas")
    add_bullets(
        doc,
        [
            "HeyGen Developer Overview: https://developers.heygen.com/",
            "HeyGen Quick Start: https://developers.heygen.com/docs/quick-start",
            "Hyperframes Overview: https://developers.heygen.com/hyperframes-overview",
            "Hyperframes API: https://developers.heygen.com/hyperframes",
            "Create Hyperframes Render reference: https://developers.heygen.com/reference/create-hyperframes-render",
            "HeyGen Usage Limits: https://developers.heygen.com/docs/usage-limits",
            "HeyGen Webhooks: https://developers.heygen.com/docs/webhooks",
            "HeyGen Pricing: https://developers.heygen.com/docs/pricing",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_h1(doc, "Apendice: checklist para decision Go/No-Go")
    add_bullets(
        doc,
        [
            "Hay pricing vigente, limites, SLA, terminos de datos y DPA aprobados.",
            "El spike demuestra que Hyperframes genera calidad aceptable para al menos un tipo de video.",
            "El output puede descargarse, validarse y almacenarse en production-videos.",
            "Los costos por video/curso son menores o justifican mejor calidad/tiempo.",
            "Existen feature flags, budgets y rollback a Remotion.",
            "Los agentes no tienen permisos de aprobacion, publicacion ni acceso directo a secretos.",
            "El equipo puede operar webhooks, logs, retries, soporte y fallos del proveedor.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
