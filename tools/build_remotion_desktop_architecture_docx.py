# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reportes/2026-07-10-remotion-desktop-architecture-analysis.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 24, 38)
MUTED = RGBColor(90, 98, 110)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN_FILL = "EAF6EF"
AMBER_FILL = "FFF6E0"
RED_FILL = "FCEAEA"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="D0D5DD", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=4)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, size=9, color=INK)
    set_table_geometry(table, widths)
    add_para(doc, "", after=4)
    return table


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        styles[style_name].font.name = "Calibri"

    header = section.header.paragraphs[0]
    header.text = "SofLIA - Engine | Decision Brief"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    paragraph_border_bottom(header, "E5E7EB", "4")
    footer = section.footer.paragraphs[0]
    footer.text = "Courseforge / Remotion Desktop Architecture"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    add_para(doc, "", after=10)
    title = add_para(doc, "Decision Memo", size=23, color=RGBColor(0, 0, 0), bold=True, after=4)
    subtitle = add_para(
        doc,
        "Separar ensamblado/renderizado Remotion hacia una app desktop",
        size=14,
        color=MUTED,
        after=14,
    )
    metadata = [
        ("Proyecto:", "Courseforge / SofLIA - Engine"),
        ("Fecha:", "10 de julio de 2026"),
        ("Rol del analisis:", "Staff Engineer / Software Architect"),
        ("Decision evaluada:", "Worker local desktop para renders pesados, con web/API como autoridad"),
        ("Recomendacion:", "Arquitectura hibrida gradual; desktop como ejecutor, cloud como fallback"),
    ]
    for label, value in metadata:
        p = add_para(doc, after=2)
        r1 = p.add_run(label + " ")
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=INK)
    rule = add_para(doc, "", after=10)
    paragraph_border_bottom(rule, "D0D5DD", "8")

    add_callout(
        doc,
        "Recomendacion ejecutiva",
        "Conviene avanzar hacia una arquitectura hibrida: la web/API conserva seguridad, estados, templates, versionado, auditoria y publicacion; una app desktop o worker local ejecuta renders pesados de Remotion con jobs firmados, tokens temporales, checksums y validacion server-side antes de aceptar el resultado.",
        GREEN_FILL,
    )

    add_heading(doc, "1. Entendimiento del sistema actual")
    add_para(
        doc,
        "El flujo actual ya separa parcialmente orquestacion y ejecucion. La UI de postproduccion selecciona template, muestra preview, inicia ensamblado y hace polling de production_jobs. La API Express autentica, valida permisos por organizacion, crea jobs idempotentes y decide el provider de render. El worker ejecuta Remotion con Chromium/FFmpeg, y Lambda/CodeBuild cubren el camino cloud para builds externos aprobados.",
    )
    add_bullet(doc, "Web: seleccion de templates, preview interno o externo, disparo de ensamblado y polling de estado.")
    add_bullet(doc, "API: autoridad de auth, tenant, permisos, idempotencia, snapshots, progress, errores y validacion de contratos.")
    add_bullet(doc, "Supabase: artifacts, material_components, production_jobs, production_assets, storage, RLS, templates, versiones y builds.")
    add_bullet(doc, "Remotion: composiciones internas en apps/web/src/remotion, renderer local en apps/api y Lambda para sites externos listos.")
    add_bullet(doc, "Templates: remotion_templates, remotion_template_versions, remotion_template_builds, estados APPROVED / APPROVED_FOR_SANDBOX / BUILT.")

    add_heading(doc, "2. Diagnostico del problema")
    add_para(
        doc,
        "Lambda y Cloud Run resuelven partes distintas del problema, pero no eliminan la tension central: renderizar video es CPU/memoria intensivo, depende de Chromium/FFmpeg, mueve assets pesados y puede durar mas que los limites practicos de un runtime serverless o administrado.",
    )
    add_simple_table(
        doc,
        ["Dominio", "Sintoma", "Lectura arquitectonica"],
        [
            ("Costo", "CPU/memoria sostenidas, min instances, almacenamiento y transferencia", "Cloud es excelente para elasticidad, pero caro como render farm permanente."),
            ("Timeout", "Renders largos o media remota lenta superan ventanas practicas", "No se debe resolver solo subiendo timeouts; hay que medir y separar fallos."),
            ("Cold start", "Chromium/bundling/FFmpeg penalizan el primer render", "Prewarm ayuda, pero no elimina renders largos."),
            ("Concurrencia", "Render secuencial o Lambda throttling", "La cola necesita leases, retries y fallback."),
            ("Debugging", "Fallos en cloud son mas opacos", "Worker local mejora reproducibilidad de errores de media, props y entorno."),
        ],
        [1600, 3000, 4760],
    )

    add_heading(doc, "3. Evaluacion de app desktop")
    add_para(
        doc,
        "Una app desktop es conveniente si se limita a ejecutar renders autorizados. No debe ser una segunda plataforma ni una autoridad paralela de templates. Su valor esta en sacar la carga pesada del cloud, permitir renders largos, controlar Chromium/FFmpeg localmente y dar mejor visibilidad operativa.",
    )
    add_heading(doc, "Ventajas", 2)
    for item in [
        "Reduce costo cloud al mover CPU/encode a maquinas locales.",
        "Permite renders largos sin depender de limites de Lambda o Cloud Run.",
        "Da control sobre versiones locales de Chromium, FFmpeg, fuentes y cache.",
        "Facilita debugging de assets, props, media remota y errores de render.",
        "Habilita trabajo offline/parcial para cache, preparacion y diagnostico.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Desventajas", 2)
    for item in [
        "Distribucion, instalacion, auto-update y soporte multi-OS.",
        "Riesgo de diferencias entre Windows, macOS y Linux.",
        "Manejo delicado de tokens, sesiones, signed URLs y permisos minimos.",
        "Uploads pesados y reanudables desde redes domesticas.",
        "Reproducibilidad: cada PC puede tener codecs, fuentes o rendimiento distinto.",
        "Mayor superficie de soporte tecnico para usuarios no tecnicos.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Alternativas comparadas")
    add_simple_table(
        doc,
        ["Alternativa", "Complejidad / costo", "Seguridad / mantenibilidad", "Veredicto"],
        [
            ("A. Todo cloud optimizando Cloud Run", "Baja-media complejidad; costo medio/alto si escala.", "Buena reproducibilidad; sigue atada a costo y timeouts cloud.", "Util para beta y fallback, no como unica apuesta."),
            ("B. Desktop como render worker", "Media complejidad; bajo costo cloud.", "Segura si usa jobs firmados y tokens temporales.", "Mejor MVP."),
            ("C. Desktop completa", "Alta complejidad; duplica UI y contratos.", "Riesgo alto de deuda y seguridad fragmentada.", "No conviene ahora."),
            ("D. Hibrida web + Supabase + desktop + cloud fallback", "Media-alta, incremental.", "Mantiene backend como autoridad y desktop como ejecutor.", "Recomendacion principal."),
            ("E. Cola competitiva local/cloud", "Alta; requiere leases, heartbeats y fairness.", "Muy flexible si se audita bien.", "Destino a mediano plazo."),
        ],
        [2100, 2300, 2500, 2460],
    )

    add_heading(doc, "5. Arquitectura recomendada")
    add_callout(
        doc,
        "Decision",
        "Implementar primero un worker local CLI sin UI, autenticado contra la API. Despues envolverlo en una app desktop con UI basica, logs, diagnostico de entorno, cache y auto-update. Cloud Run/Lambda quedan como fallback o capacidad compartida.",
        CALLOUT,
    )
    add_para(doc, "Diagrama textual del flujo:", bold=True, after=4)
    flow = [
        "Web Postproduccion",
        "  -> API /production/remotion/render",
        "  -> production_jobs PENDING",
        "  -> job firmado + provider preference local_desktop|cloud",
        "  -> Desktop Worker reclama job con lease",
        "  -> API entrega manifest autorizado: assets, templateVersionId, buildId, checksums, propsHash",
        "  -> Worker descarga assets/template aprobado",
        "  -> valida entorno local Chromium/FFmpeg/Remotion",
        "  -> renderMedia local",
        "  -> calcula checksum + duracion + logs sanitizados",
        "  -> upload a endpoint/API o signed upload URL",
        "  -> API valida resultado",
        "  -> material_components.assets.final_video_url",
        "  -> production_jobs SUCCEEDED",
        "  -> publication_requests.lesson_videos",
    ]
    for line in flow:
        p = add_para(doc, line, size=9.5, color=RGBColor(30, 41, 59), after=1)
        for run in p.runs:
            set_run_font(run, size=9.5, color=RGBColor(30, 41, 59), name="Consolas")

    add_heading(doc, "6. Stack recomendado")
    add_simple_table(
        doc,
        ["Opcion", "Uso recomendado", "Razon"],
        [
            ("Node CLI empaquetado", "MVP", "Se alinea con Remotion Renderer y la logica actual de apps/api."),
            ("Worker local sin UI", "Fase 1-2", "Menor superficie, ideal para validar contratos y seguridad."),
            ("Electron", "Fase 3", "Buen encaje con Node, logs, autoupdate, UI y control de procesos."),
            ("Tauri", "Mas adelante", "Ligero, pero exige mas puente con Node/Remotion/Chromium."),
            ("Neutralino", "No recomendado", "Demasiado limitado para render pesado con Chromium/FFmpeg."),
            ("Servicio local + UI", "Operacion madura", "Permite render en segundo plano aunque se cierre la ventana."),
        ],
        [1900, 2500, 4960],
    )

    add_heading(doc, "7. Que se mantiene en la plataforma")
    for item in [
        "Autenticacion, roles, permisos, RLS y tenant/organization_id.",
        "Artifacts, material_components, publication_requests y pipeline_events.",
        "Seleccion canonica de templates, versionado, aprobacion y build records.",
        "Estados de produccion, production_jobs, progress, errores y snapshots.",
        "Storage final, auditoria, publicacion a SofLIA y validacion server-side.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Que se mueve al desktop")
    for item in [
        "Descargar assets necesarios mediante URLs autorizadas y temporales.",
        "Resolver/cachear template aprobado o build aprobado segun manifest del backend.",
        "Ejecutar Remotion localmente con Chromium y FFmpeg controlados.",
        "Reportar progreso, logs sanitizados, errores clasificados y heartbeats.",
        "Generar MP4 final, checksum, duracion y metadata tecnica.",
        "Subir resultado via API o signed URL y esperar aceptacion del backend.",
        "Validar entorno local: version de worker, SO, CPU/RAM disponible, Chromium, FFmpeg y fuentes.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Seguridad")
    add_simple_table(
        doc,
        ["Riesgo", "Mitigacion"],
        [
            ("Templates no confiables", "Solo ejecutar versiones aprobadas; verificar bundleHash/buildHash y manifest firmado."),
            ("Tokens en desktop", "Tokens temporales por job; nunca SUPABASE_SERVICE_ROLE_KEY en cliente."),
            ("Manipulacion de jobs", "Jobs firmados, leases, worker_id, device_id y validacion de organizacion."),
            ("Uploads maliciosos", "API valida mime, tamano, checksum, duracion, job y lease antes de aceptar."),
            ("Path traversal", "Extraccion segura, allowlist de rutas, bloqueo de symlinks y rutas absolutas."),
            ("Dependencias no confiables", "Mantener allowlist, build aislado y cache por hash."),
            ("Diferencias entre PCs", "Versionar worker, Chromium, FFmpeg, fuentes y registrar capabilities."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "10. Modelo de datos y estados")
    add_para(
        doc,
        "production_jobs ya tiene la base: provider, status, snapshots, progress, provider ids, costos, duracion, errores, fechas y organizacion. La extension debe ser incremental.",
    )
    for item in [
        "Agregar provider: remotion-local-desktop, remotion-cloud-run, remotion-lambda.",
        "Agregar worker_id, device_id, claimed_at, lease_expires_at, last_heartbeat_at.",
        "Agregar worker_version, worker_platform y worker_capabilities.",
        "Agregar input_manifest_hash, asset_checksums y output_checksum.",
        "Agregar fallback_provider, retry_count y cancel_requested_at.",
        "Crear render_workers para registro, revocacion y auditoria de dispositivos.",
        "Crear render_job_events si progress jsonb deja de ser suficiente para logs.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Plan por fases")
    add_simple_table(
        doc,
        ["Fase", "Objetivo", "Resultado esperado"],
        [
            ("0", "Diagnostico y medicion de renders actuales", "Duracion, tamano de assets, errores, costo estimado y provider por job."),
            ("1", "Worker local CLI sin UI", "Un job manual renderizado localmente y aceptado por API."),
            ("2", "Integracion con cola", "Claim, lease, heartbeat, retry, cancelacion y fallback."),
            ("3", "Desktop con UI basica", "Login, estado del worker, cola, logs y diagnostico de entorno."),
            ("4", "Cache local de templates/assets", "Cache por hash, limpieza, revalidacion y reuso seguro."),
            ("5", "Operacion hibrida", "Workers locales y cloud compiten por trabajos segun politica."),
        ],
        [900, 3600, 4860],
    )

    add_heading(doc, "12. Criterios de exito del MVP")
    for item in [
        "No se modifica el contrato actual de aprobacion/versionado de templates.",
        "No se ejecutan templates no aprobados.",
        "El backend puede rechazar cualquier resultado local.",
        "Un render local completado actualiza production_jobs, material_components y publication_requests igual que cloud.",
        "El usuario no necesita Docker ni configuracion avanzada.",
        "Cloud Run/Lambda siguen disponibles como fallback.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Conclusion final")
    add_para(
        doc,
        "Conviene avanzar ahora, pero solo como arquitectura hibrida y de forma incremental. La deuda tecnica aparece si se convierte el desktop en una segunda plataforma con sus propias reglas de templates, permisos o publicacion. La linea correcta es clara: el desktop renderiza; la plataforma decide, valida y audita.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT.resolve())
