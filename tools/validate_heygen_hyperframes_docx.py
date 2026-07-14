from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "reportes" / "2026-07-10-heygen-hyperframes-architecture-analysis.docx"


required_phrases = [
    "Recomendacion ejecutiva",
    "Entendimiento del sistema actual",
    "Diagnostico del problema actual",
    "Evaluacion de HeyGen Hyperframes",
    "Arquitectura recomendada",
    "Costos",
    "Seguridad, privacidad y compliance",
    "Plan de implementacion por fases",
    "Recomendacion final",
]


def main() -> None:
    doc = Document(DOCX)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    text = "\n".join(parts)
    missing = [phrase for phrase in required_phrases if phrase not in text]
    if missing:
        raise SystemExit("Missing required phrases: " + ", ".join(missing))
    print(f"docx={DOCX}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"sections={len(doc.sections)}")


if __name__ == "__main__":
    main()
